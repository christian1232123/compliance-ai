import os
import json
import sqlite3
from io import BytesIO
from typing import Optional
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from fastapi.requests import Request
from pypdf import PdfReader
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google import genai
from google.genai import types

# Inizializzazione FastAPI
app = FastAPI()

# Database SQLite Integrato
DB_PATH = "audit_history.db"

def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS audit_reports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT,
            standard TEXT,
            language TEXT,
            risk_score INTEGER,
            risk_level TEXT,
            summary TEXT,
            report_markdown TEXT,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()
    conn.close()

init_db()

# Inizializzazione Client Gemini
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")
client = genai.Client(api_key=GEMINI_API_KEY) if GEMINI_API_KEY else None

STRIPE_SECRET_KEY = os.environ.get("STRIPE_SECRET_KEY", "sk_test_mock")

@app.get("/", response_class=HTMLResponse)
async def read_root():
    with open("index.html", "r", encoding="utf-8") as f:
        return f.read()

@app.post("/analyze")
async def analyze_pdf(
    file: UploadFile = File(...),
    standard: str = Form("gdpr"),
    language: str = Form("it")
):
    if not file.filename.endswith('.pdf'):
        raise HTTPException(status_code=400, detail="Il file deve essere un PDF.")

    try:
        # Lettura PDF
        pdf_bytes = await file.read()
        reader = PdfReader(BytesIO(pdf_bytes))
        extracted_text = ""
        for page in reader.pages:
            text = page.extract_text()
            if text:
                extracted_text += text + "\n"

        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="Impossibile estrarre testo dal PDF.")

        # Tronca testo se troppo lungo
        extracted_text = extracted_text[:15000]

        if not client:
            raise HTTPException(status_code=500, detail="Chiave API Gemini non configurata.")

        # Mappatura lingue
        lang_map = {
            "it": "Italiano",
            "en": "English",
            "es": "Español",
            "de": "Deutsch"
        }
        target_lang = lang_map.get(language, "Italiano")

        # Mappatura standard
        std_map = {
            "gdpr": "GDPR & Normativa Privacy europea",
            "iso27001": "ISO/IEC 27001 (Sicurezza delle Informazioni)",
            "sicurezza": "D.Lgs 81/08 (Sicurezza sul Lavoro)"
        }
        target_std = std_map.get(standard, "GDPR")

        prompt = f"""
        Sei un Auditor di Compliance esperto. Analizza il seguente testo estratto da un documento PDF ed effettua una valutazione di conformità rispetto allo standard: {target_std}.
        Rispondi ESCLUSIVAMENTE nella lingua: {target_lang}.

        Restituisci la risposta IN FORMATO JSON STRETTO con le seguenti chiavi:
        1. "risk_score": un numero intero da 0 a 100 dove 100 indica elevata conformità / basso rischio, e 0 indica totale non conformità / alto rischio.
        2. "risk_level": stringa tra "Basso", "Medio", "Alto", "Critico".
        3. "summary": una breve sintesi di 2-3 frasi sull'esito globale dell'audit.
        4. "markdown_report": la relazione completa dettagliata formattata in Markdown (include sezioni: Punti di Forza, Criticità/Non Conformità, Rischi Legali/Privacy, Raccomandazioni Operative).

        Testo del documento:
        {extracted_text}
        """

        response = client.models.generate_content(
            model='gemini-2.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type="application/json"
            )
        )

        result_data = json.loads(response.text)

        # Salvataggio nel Database
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO audit_reports (filename, standard, language, risk_score, risk_level, summary, report_markdown)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        ''', (
            file.filename,
            target_std,
            target_lang,
            result_data.get("risk_score", 70),
            result_data.get("risk_level", "Medio"),
            result_data.get("summary", ""),
            result_data.get("markdown_report", "")
        ))
        report_id = cursor.lastrowid
        conn.commit()
        conn.close()

        result_data["report_id"] = report_id
        return JSONResponse(content=result_data)

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/history")
async def get_history():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT id, filename, standard, language, risk_score, risk_level, summary, created_at FROM audit_reports ORDER BY created_at DESC LIMIT 20')
    rows = cursor.fetchall()
    conn.close()
    return JSONResponse(content=[dict(row) for row in rows])

@app.get("/get-report/{report_id}")
async def get_report_by_id(report_id: int):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM audit_reports WHERE id = ?', (report_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Report non trovato")
    return JSONResponse(content=dict(row))

@app.get("/export-docx/{report_id}")
async def export_docx(report_id: int):
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM audit_reports WHERE id = ?', (report_id,))
    row = cursor.fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=404, detail="Report non trovato")

    doc = Document()
    
    # Titolo
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Report di Audit & Compliance")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(16, 185, 129)

    # Info
    doc.add_paragraph(f"Documento: {row['filename']}")
    doc.add_paragraph(f"Standard: {row['standard']}")
    doc.add_paragraph(f"Punteggio di Conformità: {row['risk_score']}/100 (Livello di Rischio: {row['risk_level']})")
    doc.add_paragraph(f"Data: {row['created_at']}")
    doc.add_heading("Sintesi", level=2)
    doc.add_paragraph(row['summary'])

    doc.add_heading("Dettaglio Audit", level=2)
    doc.add_paragraph(row['report_markdown'])

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename_out = f"Audit_Report_{report_id}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename_out}"}
    )

@app.post("/create-checkout-session")
async def create_checkout_session(plan: str = Form(...)):
    import stripe
    stripe.api_key = STRIPE_SECRET_KEY
    price = 4900 if plan == "pro" else 19900
    try:
        session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{
                'price_data': {
                    'currency': 'eur',
                    'product_data': {'name': f'Piano ComplianceAI {plan.capitalize()}'},
                    'unit_amount': price,
                },
                'quantity': 1,
            }],
            mode='payment',
            success_url='https://compliance-ai-qx5a.onrender.com/?success=true',
            cancel_url='https://compliance-ai-qx5a.onrender.com/?canceled=true',
        )
        return {"url": session.url}
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})
